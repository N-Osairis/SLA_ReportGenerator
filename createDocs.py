from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from docx.shared import Mm
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENTATION
from docx.enum.section import WD_SECTION_START
import lxml
from docx.shared import RGBColor
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

import calendar
import datetime

# this class provided to create the report in microsoft word format based on the input argument
class creatWord:

    def __init__(self, Month1, Year1, table_list, outputadd, clientname, graph_no, SLA, RD):
        self.month =datetime.date(1900, Month1, 1).strftime('%B')
        self.year = Year1
        self.lastday = calendar.monthrange(Year1, Month1)[1]
        self.table_list = table_list
        self.outputadd = outputadd
        self.cliname = clientname
        self.NoOfgraph = graph_no
        self.slaValue = SLA
        self.rd = RD

    # this function create the word file based on the requested fromat
    def creatingfunction_word(self):
        # word setup page-----------------------
        # creating document file
        document = Document()
        # creating 2 section for separate layout
        section = document.sections[0]
        footer = section.footer
        paragraph = footer.paragraphs[0]
        # if "PMP"  in self.cliname:
        # paragraph.text = "PAN MALAYSIAN POOL SLA REPORT"
        # elif  "FJB"  in self.cliname:
        paragraph.text = "N’osairis Technology Solutions Sdn Bhd \n " \
                         "Unit 9-6, Level 9, Tower B, Vertical Business Suite 2, " \
                         "Avenue 3, Bangsar South, No. 8, Jalan Kerinchi, 59200 Kuala Lumpur."
        paragraph.style = document.styles["Footer"]
        paragraph.paragraph_format.left_indent = Inches(-0.5)
        paragraph.paragraph_format.right_indent = Inches(-0.5)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        paragraph.underline = True
        pr=paragraph.runs[0]
        pr.font.size = Pt(8)
        pr.font.name = 'Arial'
        pr.bold = True
        # creating cover page------------------------------------------
        H1 = document.add_paragraph()
        H1.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        H1_run = H1.add_run('N’OSAIRIS TECHNOLOGY SOLUTIONS'
                            '                CONFIDENTIAL')
        H1_run.bold = True
        H1_run.font.size = Pt(20)
        H1_run.font.name = 'Calibri (Body)'
        H1_format = H1.paragraph_format
        H1_format.left_indent = Inches(-0.6)
        H1_format.right_indent = Inches(-0.6)

        line = document.add_picture('{}\\line.png'.format(self.rd), width=Inches(7))
        last_paragraph = document.paragraphs[-1]
        last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        last_paragraph_format = last_paragraph.paragraph_format
        last_paragraph_format.left_indent = Inches(-0.6)
        last_paragraph_format.right_indent = Inches(-0.6)

        document.add_paragraph("\n")
        if "PMP" in self.cliname:
            logo = document.add_picture('{}\\{}.png'.format(self.rd, self.cliname),
                                        width=Inches(4.7), height=Inches(1.1))
        elif "FJB" in self.cliname:
            logo = document.add_picture('{}\\{}.png'.format(self.rd,self.cliname),
                                      width=Inches(4.7),height=Inches(1.1))
        elif "VINX-EMONEY" in self.cliname or "VINX-WAN" in self.cliname:
            logo = document.add_picture('{}\\{}.png'.format(self.rd, self.cliname),
                                        width=Inches(3.1), height=Inches(1.6))

        last_paragraph = document.paragraphs[-1]
        last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        last_paragraph_format = last_paragraph.paragraph_format
        last_paragraph_format.left_indent = Inches(-0.7)
        last_paragraph_format.right_indent = Inches(-0.3)
        document.add_paragraph("\n")
        p1 = document.add_paragraph()
        p1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        if "PMP" in self.cliname:
            P1_run = p1.add_run('PAN MALAYSIAN POOL AVAILABILITY REPORT')
        elif "FJB" in self.cliname:
            P1_run = p1.add_run('FJ Benjamin Holdings Ltd')
        elif "VINX-EMONEY" in self.cliname or "VINX-WAN" in self.cliname:
            P1_run = p1.add_run('VINX MALAYSIA SDN. BHD')
        P1_run.bold = True
        P1_run.font.size=Pt(18)
        P1_run.font.name = 'Calibri (Body)'
        document.add_paragraph("\n\n")
        p1_1 = document.add_paragraph()
        p1_1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p1_1_run = p1_1 .add_run('AVAILABILITY REPORT')
        p1_1_run.bold = True
        p1_1_run.font.size=Pt(18)
        p1_1_run.font.name = 'Calibri (Body)'
        p2 = document.add_paragraph()
        p2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        P2_run = p2.add_run('1 {} {} - {} {} {} '.format(self.month, self.year, self.lastday, self.month, self.year ))
        P2_run.bold = True
        P2_run.font.size = Pt(18)
        P2_run.font.name = 'Calibri (Body)'
        document.add_paragraph("\n \n \n")
        p3 = document.add_paragraph()
        p3.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        P3_run = p3.add_run('Submitted by')
        P3_run.bold = False
        P3_run.font.size = Pt(14)
        P3_run.font.name = 'Calibri (Body)'
        # document.add_paragraph("\n")
        document.add_picture('{}\\Nosairis.png'.format(self.rd), width=Inches(1.49), height=Inches(1.8))
        last_paragraph = document.paragraphs[-1]
        last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # creating section breack ------------------------------
        section = document.sections[0]
        section.page_height = Mm(297)
        section.page_width = Mm(210)
        section.start_type = WD_SECTION_START.NEW_PAGE
        new_section = document.add_section(WD_SECTION_START.NEW_PAGE)
        print('len(sections) = {}'.format(len(document.sections)))
        new_section.orientation = WD_ORIENTATION.LANDSCAPE
        new_width, new_height = new_section.page_height, new_section.page_width
        new_section.page_width = new_width
        new_section.page_height = new_height
        # =============================================
        pt = document.add_paragraph()
        pt.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        Pt_run = pt.add_run('AVAILABILITY REPORT :')
        Pt_run.bold = True
        Pt_run.font.size = Pt(14)
        Pt_run.font.name = 'Calibri (Body)'
        pt_format = pt.paragraph_format
        pt_format.left_indent = Inches(-0.5)
        pt_format.right_indent = Inches(-0.5)

        pt2 = document.add_paragraph()
        pt2.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        Pt2_run = pt2.add_run('Date :')
        Pt2_run.bold = True
        Pt2_run.font.size = Pt(14)
        Pt2_run.font.name = 'Calibri (Body)'
        pt2_format = pt2.paragraph_format
        pt2_format.left_indent = Inches(-0.5)
        pt2_format.right_indent = Inches(-0.5)

        ptd = document.add_paragraph()
        ptd.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        Ptd_run = pt2.add_run('1 {} {} - {} {} {} '.format(self.month, self.year, self.lastday, self.month, self.year))
        Ptd_run.bold = False
        Ptd_run.font.size = Pt(12)
        Ptd_run.font.name = 'Calibri (Body)'
        ptd_format = ptd.paragraph_format
        ptd_format.left_indent = Inches(-0.5)
        ptd_format.right_indent = Inches(-0.5)

        # creating 2nd page (SLA Tables)
        ptsla = document.add_paragraph()
        ptsla.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        ptsla_run = ptsla.add_run('SLA Table:')
        ptsla_run.bold = True
        ptsla_run.Strickthrough = True
        ptsla_run.font.size = Pt(14)
        ptsla_run.font.name = 'Calibri (Body)'
        ptsla_format = ptsla.paragraph_format
        ptsla_format.left_indent = Inches(-0.5)
        ptsla_format.right_indent = Inches(-0.5)
        # setting the SLA row Tables)
        if "VINX-WAN" in self.cliname:
            table = document.add_table(rows=1, cols=8)
            table_head = ['NO', 'Site Name', 'Connectivity', 'Business Hours \n Minutes', 'WAN \n Downtime Minutes', 'Customer \n Downtime Minutes',
                         'Total Downtime Minutes\n (WAN+Customer)', 'WAN\nAvailability\nPercentage']

        else:
            table = document.add_table(rows=1, cols=9)
            table_head = ['NO', 'Site Name', 'Connectivity', 'From Date', 'To Date', 'Exact Site Uptime (mins)',
                         'Total Availability Time (mins)', 'Total downtime', 'Uptime (%)']

        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.allow_autofit = False
        # ======================================================
        nsmap = table._element[0].nsmap  # For namespaces
        searchtag = '{%s}tblPr' % nsmap['w']  # w:tblPr
        mytag = '{%s}tblInd' % nsmap['w']  # w:tblInd
        myw = '{%s}w' % nsmap['w']  # w:w
        mytype = '{%s}type' % nsmap['w']  # w:type
        for elt in table._element:
            if elt.tag == searchtag:
                myelt = lxml.etree.Element(mytag)
                myelt.set(myw, '-1200')
                myelt.set(mytype, 'dxa')
                myelt = elt.append(myelt)
        #=============================================
        hdr_cells = table.rows[0].cells
        # tablehead=list(self.table_list.columns)
        if "VINX-WAN" in self.cliname:
            y=0
            while y<=7:
                # if y<2:
                #  hdr_cells[y].text = tablehead[y]
                # else:
                #  hdr_cells[y].text = tablehead[y+1]
                hdr_cells[y].text = table_head[y]
                shading_elm = parse_xml(r'<w:shd {} w:fill="e6f026"/>'.format(nsdecls('w')))
                hdr_cells[y]._tc.get_or_add_tcPr().append(shading_elm)
                hdr_cells[y]._tc.tcPr.tcW.type = 'auto'
                paragraph = hdr_cells[y].paragraphs[0]
                paragraph.alignment =WD_PARAGRAPH_ALIGNMENT.CENTER
                run = paragraph.runs
                font = run[0].font
                font.name = 'Calibri (Body)'
                font.size = Pt(9)  # font size = 30
                font.bold = True
                y += 1
            count = 1
            for x in self.table_list.sort_values(by='Site Name').index:
                # 'NO', 'Site Name', 'Connectivity', 'From Date', 'To Date',
                #                                                   'Exact Site Uptime (mins)',
                #                                                   'Total Availability Time-Business Hours (mins)',
                #                                                   'Total downtime_WAN', 'Customer Downtime (mins)',
                #                                                   'Total Downtime(wan+customer)(mins)', 'Uptime (%)
                row_cells = table.add_row().cells
                row_cells[0].text = str(count)
                row_cells[1].text = str(self.table_list['Site Name'][x])
                # row_cells[2].text = str(self.table_list['TNS Router ID'][x])
                row_cells[2].text = str(self.table_list['Connectivity'][x])
                row_cells[3].text = str(self.table_list['Total Availability Time-Business Hours (mins)'][x])
                row_cells[4].text = str(self.table_list['Total downtime_WAN'][x])
                # row_cells[5].text = str(self.table_list['Exact Site Uptime (mins)'][x])
                row_cells[5].text = str(self.table_list['Customer Downtime (mins)'][x])
                row_cells[6].text = str(self.table_list['Total Downtime(wan+customer)(mins)'][x])
                row_cells[7].text = str(self.table_list['Uptime (%)'][x])

                count += 1
                y = 0
                while y <= 7:
                    paragraph = row_cells[y].paragraphs[0]
                    row_cells[y]._tc.tcPr.tcW.type = 'auto'
                    if y == 1 or y == 2:
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    else:
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    run = paragraph.runs
                    font = run[0].font
                    font.size = Pt(9)  # font size = 30
                    font.name = 'Calibri (Body)'
                    if float(self.table_list['Uptime (%)'][x])<self.slaValue:
                        font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
                    else:
                        font.color.rgb = RGBColor(0x00, 0x00, 0x00)
                    y += 1

        else :
            y = 0
            while y <= 8:
                # if y<2:
                #  hdr_cells[y].text = tablehead[y]
                # else:
                #  hdr_cells[y].text = tablehead[y+1]
                hdr_cells[y].text = table_head[y]
                shading_elm = parse_xml(r'<w:shd {} w:fill="e6f026"/>'.format(nsdecls('w')))
                hdr_cells[y]._tc.get_or_add_tcPr().append(shading_elm)
                hdr_cells[y]._tc.tcPr.tcW.type = 'auto'
                paragraph = hdr_cells[y].paragraphs[0]
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = paragraph.runs
                font = run[0].font
                font.name = 'Calibri (Body)'
                font.size = Pt(9)  # font size = 30
                font.bold = True
                y += 1
            count = 1
            for x in self.table_list.sort_values(by='Site Name').index:
                # 'NO', 'Site Name', 'Connectivity', 'From Date', 'To Date',
                #                                                   'Exact Site Uptime (mins)',
                #                                                   'Total Availability Time-Business Hours (mins)',
                #                                                   'Total downtime_WAN', 'Customer Downtime (mins)',
                #                                                   'Total Downtime(wan+customer)(mins)', 'Uptime (%)
                row_cells = table.add_row().cells
                row_cells[0].text = str(count)
                row_cells[1].text = str(self.table_list['Site Name'][x])
                # row_cells[2].text = str(self.table_list['TNS Router ID'][x])
                row_cells[2].text = str(self.table_list['Connectivity'][x])
                row_cells[3].text = str(self.table_list['From Date'][x])
                row_cells[4].text = str(self.table_list['To Date'][x])
                row_cells[5].text = str(self.table_list['Exact Site Uptime (mins)'][x])
                row_cells[6].text = str(self.table_list['Total Availability Time-Business Hours (mins)'][x])
                row_cells[7].text = str(self.table_list['Total downtime_WAN'][x])
                row_cells[8].text = str(self.table_list['Uptime (%)'][x])
                count += 1
                y = 0
                while y <= 8:
                    paragraph = row_cells[y].paragraphs[0]
                    row_cells[y]._tc.tcPr.tcW.type = 'auto'
                    if y == 1 or y == 2:
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    else:
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    run = paragraph.runs
                    font = run[0].font
                    font.size = Pt(9)  # font size = 30
                    font.name = 'Calibri (Body)'
                    if float(self.table_list['Uptime (%)'][x]) < self.slaValue:
                        font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
                    else:
                        font.color.rgb = RGBColor(0x00, 0x00, 0x00)
                    y += 1
        #creating new section to add the graphs
        document.add_page_break()
        p5 = document.add_paragraph()
        p5.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        p5_run = p5.add_run('SLA Graphs :')
        p5_run.bold = True
        p5_run.font.size = Pt(14)
        p5_format = p5.paragraph_format
        p5_format.left_indent = Inches(-0.5)
        p5_format.right_indent = Inches(-0.5)

        img_count = 1
        while img_count <= self.NoOfgraph:
            document.add_picture('{}\\{}.png'.format(self.outputadd, img_count), width=Inches(8.2), height=Inches(2.8))
            last_paragraph = document.paragraphs[-1]
            last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            last_paragraph_format = last_paragraph.paragraph_format
            last_paragraph_format.left_indent = Inches(-0.5)
            last_paragraph_format.right_indent = Inches(-0.5)

            img_count += 1
        document.save('{}\\{}-{}_SLA_Report.docx'.format(self.outputadd, self.cliname, self.month) )